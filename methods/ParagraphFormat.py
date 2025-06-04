from docx import Document
from docx.shared import Cm, RGBColor
import json

def open_base_rules(file: str) -> dict:
    """
    Функция для открытия файла с правилами
    :param file:
    :return:
    """
    with open(file, 'r', encoding='utf-8') as file:
        data = json.load(file)
        base_rules = data["base_rules"]
        return base_rules


def checking_paragraphs_on_line_spacing(base_rules: dict, doc: Document) -> dict:
    """
    Функция для проверки всех абзацев на соответствие заданным шрифтам и прочим параметрам
    :param base_rules:
    :param doc:
    :return:
    """
    number_paragraph = 1
    paragraph_status = {
        "Общее количество абзацев": len(doc.paragraphs),
        "Количество абзацев, соответствующих правилам": 0,
        "Нетекстовые абзацы": 0,
        "Количество абзацев с нарушениями": 0,
        "Абзацы с нарушениями": {

        }
    }

    def error_paragraph_adding(para, error_type):
        """
        Добавление информации об абзаце с нарушениями форматирования
        """
        paragraph_status["Количество абзацев с нарушениями"] += 1
        paragraph_status["Абзацы с нарушениями"][number_paragraph] = {
            "Нарушенный стиль": para.style.name,
            "Номер абзаца": number_paragraph,
            "Часть текста абзаца": para.text,
            "Тип ошибки": error_type
        }

    alignment_dict = {
        1: "По центру",
        2: "По правому краю",
        3: "По ширине"
    }



    black_rgb = RGBColor(0x00, 0x00, 0x00)  # код черного цвета для сравнения
    paragraphs = doc.paragraphs
    for paragraph in paragraphs:
        # Проверяет, является ли абзац текстовым или же нет
        if paragraph._p.xpath(".//w:drawing") or paragraph._p.xpath(".//w:tbl") or paragraph._p.xpath(".//w:object"):
            paragraph_status["Нетекстовые абзацы"] += 1
            continue
        paragraph_style = paragraph.style
        if paragraph_style.name in base_rules:
            paragraph_format = paragraph.paragraph_format
            style_paragraph_name =  paragraph_style.name # название используемого стиля
            paragraph_style_font = paragraph_style.font.name
            paragraph_color = "Черный" if paragraph_style.font.color.rgb == black_rgb or None else "Другой"
            paragraph_bold = "Да" if paragraph_style.font.bold else "Нет"
            paragraph_italic = "Да" if paragraph_style.font.italic else "Нет"
            paragraph_underline = "Да" if paragraph_style.font.underline else "Нет"
            paragraph_size = paragraph_style.font.size.pt if paragraph_style.font.size else None

            # если у абзаца используется межстрочный интервал, который не наследуется от используемого стиля
            if paragraph_format.line_spacing:
                paragraph_spacing = paragraph_format.line_spacing
            else:
                if paragraph_style.paragraph_format.line_spacing:
                    paragraph_spacing = paragraph_style.paragraph_format.line_spacing
                else:
                    paragraph_spacing = 0.0

            # если у абзаца есть отступ первой строки, который не наследуется от используемого стиля
            if paragraph_format.first_line_indent:
                first_line_indent = paragraph_format.first_line_indent.cm
            else:
                # если используется отступ от родительского стиля
                if paragraph_style.paragraph_format.first_line_indent:
                    first_line_indent = paragraph_style.paragraph_format.first_line_indent.cm
                # если отступ у родительского стиля не определен
                else:
                    first_line_indent = 0.0

            # если у абзаца есть межстрочный отступ после абзаца
            if paragraph_format.space_after:
                space_after = paragraph_format.space_after.pt
            else:
                # если используется межстрочный отступ после абзаца от родительского стиля
                if paragraph_style.paragraph_format.space_after:
                    space_after = paragraph_style.paragraph_format.space_after.pt
                else:
                    space_after = 0.0

            # если у абзаца есть межстрочный отступ до абзаца
            if paragraph_format.space_before:
                space_before = paragraph_format.space_before.pt
            else:
                # если используется межстрочный отступ до абзаца от родительского стиля
                if paragraph_style.paragraph_format.space_before:
                    space_before = paragraph_style.paragraph_format.space_before.pt
                else:
                    space_before = 0.0

            # если у абзаца изменено выравнивание
            if paragraph_format.alignment:
                alignment = alignment_dict[paragraph_format.alignment]
            else:
                if paragraph_style.paragraph_format.alignment:
                    alignment = alignment_dict[paragraph_style.paragraph_format.alignment]
                else:
                    alignment = "По левому краю"

            paragraph_style_status = {"Используемый шрифт": paragraph_style_font,
                                      "Размер шрифта": paragraph_size,
                                      "Цвет шрифта": paragraph_color,
                                      "Жирный": paragraph_bold,
                                      "Курсивный": paragraph_italic,
                                      "Подчеркнутый": paragraph_underline,
                                      "Межстрочный интервал": float(paragraph_spacing),
                                      "Отступ первой строки": float(round(first_line_indent, 2)),
                                      "Отступ после абзаца": float(space_after),
                                      "Отступ перед абзацем": float(space_before),
                                      "Выравнивание": alignment
            }
            # print(f"{number_paragraph}, {style_paragraph_name}; {paragraph_style_font}; {paragraph_color};"
            #       f"{paragraph_bold}; {paragraph_italic}; {paragraph_underline};"
            #       f"{paragraph_size}; {paragraph_spacing}; {first_line_indent};"
            #       f"{space_after}; {space_before}; {alignment}")
            # if number_paragraph == 275 or number_paragraph == 377:
            #     print(paragraph.text)
            if paragraph_style_status == base_rules[style_paragraph_name]:
                paragraph_status["Количество абзацев, соответствующих правилам"] += 1
            else:
                error_paragraph_adding(paragraph, "Стиль абзаца отличается от родительского")
        else:
            error_paragraph_adding(paragraph, "Стиль абзаца не из базы правил")
        number_paragraph += 1
    return paragraph_status

# def main(path: str):
#     try:
#         document = Document(path)
#         check_result = checking_paragraphs_on_line_spacing(open_base_rules("../rules/font_rules.json") ,document)
#         dictionary_output(check_result)
#         with open('../paragraph_result.json', 'w', encoding='utf-8') as f:
#             json.dump(check_result, f, ensure_ascii = False, indent = 1)
#     except Exception as e:
#         print(f"Ошибка: {e}")
#
# if __name__ == "__main__":
#     main("C:\\Users\\danil\\Desktop\\Lectures\\DiplomWork\\ВКР, текст.docx")